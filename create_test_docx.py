"""
테스트용 DOCX 파일 생성 스크립트
글 정렬, 폰트 색깔, 들여쓰기 기능을 테스트합니다.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_docx():
    doc = Document()

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("한글 파일 변환 기능 테스트")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0, 0, 255)  # 파란색

    doc.add_paragraph()  # 빈 줄

    # 1. 폰트 색깔 테스트
    section1 = doc.add_paragraph()
    section1_run = section1.add_run("1. 폰트 색깔 테스트")
    section1_run.bold = True
    section1_run.font.size = Pt(14)

    # 빨간색 텍스트
    p1 = doc.add_paragraph()
    run1 = p1.add_run("이 텍스트는 빨간색입니다.")
    run1.font.color.rgb = RGBColor(255, 0, 0)
    run1.font.size = Pt(12)

    # 초록색 텍스트
    p2 = doc.add_paragraph()
    run2 = p2.add_run("이 텍스트는 초록색입니다.")
    run2.font.color.rgb = RGBColor(0, 255, 0)
    run2.font.size = Pt(12)

    # 파란색 텍스트
    p3 = doc.add_paragraph()
    run3 = p3.add_run("이 텍스트는 파란색입니다.")
    run3.font.color.rgb = RGBColor(0, 0, 255)
    run3.font.size = Pt(12)

    # 보라색 텍스트
    p4 = doc.add_paragraph()
    run4 = p4.add_run("이 텍스트는 보라색입니다.")
    run4.font.color.rgb = RGBColor(128, 0, 128)
    run4.font.size = Pt(12)

    # 주황색 텍스트
    p5 = doc.add_paragraph()
    run5 = p5.add_run("이 텍스트는 주황색입니다.")
    run5.font.color.rgb = RGBColor(255, 165, 0)
    run5.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 2. 글 정렬 테스트
    section2 = doc.add_paragraph()
    section2_run = section2.add_run("2. 글 정렬 테스트")
    section2_run.bold = True
    section2_run.font.size = Pt(14)

    # 왼쪽 정렬
    p_left = doc.add_paragraph("이 텍스트는 왼쪽 정렬입니다. (LEFT)")
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_left.runs[0].font.size = Pt(12)

    # 가운데 정렬
    p_center = doc.add_paragraph("이 텍스트는 가운데 정렬입니다. (CENTER)")
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_center.runs[0].font.size = Pt(12)

    # 오른쪽 정렬
    p_right = doc.add_paragraph("이 텍스트는 오른쪽 정렬입니다. (RIGHT)")
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.runs[0].font.size = Pt(12)

    # 양쪽 정렬
    p_justify = doc.add_paragraph(
        "이 텍스트는 양쪽 정렬입니다. 양쪽 정렬은 텍스트가 길 때 더 명확히 보입니다. "
        "양쪽 정렬을 사용하면 텍스트가 양쪽 끝에 맞춰져 깔끔하게 보입니다. (JUSTIFY)"
    )
    p_justify.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_justify.runs[0].font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 3. 들여쓰기 테스트
    section3 = doc.add_paragraph()
    section3_run = section3.add_run("3. 들여쓰기 테스트")
    section3_run.bold = True
    section3_run.font.size = Pt(14)

    # 첫줄 들여쓰기
    p_first_indent = doc.add_paragraph(
        "이 문단은 첫줄 들여쓰기가 적용되어 있습니다. "
        "첫 번째 줄만 들여쓰기가 되고 나머지 줄은 정상적으로 표시됩니다. "
        "문단의 시작을 명확히 하는데 유용합니다."
    )
    p_first_indent.paragraph_format.first_line_indent = Inches(0.5)
    p_first_indent.runs[0].font.size = Pt(12)

    # 왼쪽 들여쓰기
    p_left_indent = doc.add_paragraph(
        "이 문단은 왼쪽 들여쓰기가 적용되어 있습니다. "
        "모든 줄이 왼쪽에서 들여써집니다. "
        "인용문이나 하위 항목을 표시할 때 유용합니다."
    )
    p_left_indent.paragraph_format.left_indent = Inches(0.5)
    p_left_indent.runs[0].font.size = Pt(12)

    # 왼쪽 + 첫줄 들여쓰기
    p_combined_indent = doc.add_paragraph(
        "이 문단은 왼쪽 들여쓰기와 첫줄 들여쓰기가 모두 적용되어 있습니다. "
        "전체가 왼쪽에서 들여써지고, 첫 줄은 추가로 더 들여써집니다. "
        "구조화된 문서에서 계층을 표현할 때 유용합니다."
    )
    p_combined_indent.paragraph_format.left_indent = Inches(0.5)
    p_combined_indent.paragraph_format.first_line_indent = Inches(0.3)
    p_combined_indent.runs[0].font.size = Pt(12)

    # 오른쪽 들여쓰기
    p_right_indent = doc.add_paragraph(
        "이 문단은 오른쪽 들여쓰기가 적용되어 있습니다. "
        "오른쪽 여백이 줄어들어 텍스트 영역이 좁아집니다. "
        "특별한 강조가 필요한 문단에 사용할 수 있습니다."
    )
    p_right_indent.paragraph_format.right_indent = Inches(0.5)
    p_right_indent.runs[0].font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 4. 내어쓰기 테스트
    section4 = doc.add_paragraph()
    section4_run = section4.add_run("4. 내어쓰기(Hanging Indent) 테스트")
    section4_run.bold = True
    section4_run.font.size = Pt(14)

    p_hanging_indent = doc.add_paragraph(
        "이 문단은 내어쓰기가 적용되어 있습니다. 첫 줄은 그대로이고, "
        "두 번째 줄부터 들여쓰기가 되어 목록이나 인용에 자주 사용됩니다. "
        "글머리 기호 뒤에 오는 텍스트를 정렬할 때 유용합니다."
    )
    p_hanging_indent.paragraph_format.left_indent = Inches(0.5)
    p_hanging_indent.paragraph_format.first_line_indent = Inches(-0.25)
    p_hanging_indent.runs[0].font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 5. 복합 테스트 (색깔 + 정렬 + 들여쓰기)
    section5 = doc.add_paragraph()
    section5_run = section5.add_run("5. 복합 서식 테스트")
    section5_run.bold = True
    section5_run.font.size = Pt(14)

    # 빨간색 + 가운데 정렬
    p_red_center = doc.add_paragraph()
    p_red_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_red_center = p_red_center.add_run("빨간색 + 가운데 정렬")
    run_red_center.font.color.rgb = RGBColor(255, 0, 0)
    run_red_center.font.size = Pt(12)
    run_red_center.bold = True

    # 파란색 + 오른쪽 정렬 + 들여쓰기
    p_blue_right = doc.add_paragraph()
    p_blue_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_blue_right.paragraph_format.right_indent = Inches(0.5)
    run_blue_right = p_blue_right.add_run("파란색 + 오른쪽 정렬 + 오른쪽 들여쓰기")
    run_blue_right.font.color.rgb = RGBColor(0, 0, 255)
    run_blue_right.font.size = Pt(12)
    run_blue_right.italic = True

    # 초록색 + 왼쪽 정렬 + 들여쓰기
    p_green_left = doc.add_paragraph()
    p_green_left.paragraph_format.left_indent = Inches(1.0)
    p_green_left.paragraph_format.first_line_indent = Inches(0.3)
    run_green_left = p_green_left.add_run(
        "초록색 + 왼쪽 들여쓰기 + 첫줄 들여쓰기. "
        "여러 서식을 동시에 적용하여 복잡한 문서를 만들 수 있습니다."
    )
    run_green_left.font.color.rgb = RGBColor(0, 128, 0)
    run_green_left.font.size = Pt(12)

    # 파일 저장
    output_path = "test_formatting.docx"
    doc.save(output_path)
    print(f"[OK] 테스트 파일 생성 완료: {output_path}")
    print("\n생성된 테스트 항목:")
    print("  - 폰트 색깔: 빨강, 초록, 파랑, 보라, 주황")
    print("  - 글 정렬: 왼쪽, 가운데, 오른쪽, 양쪽정렬")
    print("  - 들여쓰기: 첫줄, 왼쪽, 오른쪽, 복합")
    print("  - 내어쓰기: Hanging Indent")
    print("  - 복합 서식: 색깔+정렬+들여쓰기 조합")

    return output_path

if __name__ == "__main__":
    create_test_docx()
