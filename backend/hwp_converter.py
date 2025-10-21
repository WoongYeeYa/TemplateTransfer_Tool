"""
한글 2018 COM API를 이용한 DOCX to HWP 변환 모듈
완전 수동 재구성 버전 - 모든 요소를 COM API로 직접 그림
"""
import os
import win32com.client
import pythoncom
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
import logging

# 로깅 설정
log_formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
log_handler = logging.FileHandler("converter.log", mode='w', encoding='utf-8')
log_handler.setFormatter(log_formatter)
root_logger = logging.getLogger()
if root_logger.hasHandlers():
    root_logger.handlers.clear()
root_logger.addHandler(log_handler)
root_logger.setLevel(logging.DEBUG)

logger = logging.getLogger(__name__)

# -----------------------------------------------------------------------------
# 한글 2018 COM API 헬퍼 함수들
# -----------------------------------------------------------------------------

def _hwp_insert_text(hwp, text):
    """텍스트를 HWP에 삽입"""
    if not text:
        return
    try:
        hwp.HAction.GetDefault("InsertText", hwp.HParameterSet.HInsertText.HSet)
        hwp.HParameterSet.HInsertText.Text = text
        hwp.HAction.Execute("InsertText", hwp.HParameterSet.HInsertText.HSet)
    except Exception as e:
        logger.error(f"텍스트 삽입 실패: {e}")

def _hwp_break_para(hwp):
    """문단 나누기"""
    try:
        hwp.HAction.Run("BreakPara")
    except Exception as e:
        logger.error(f"문단 나누기 실패: {e}")

def _hwp_set_char_shape(hwp, bold=None, italic=None, underline=None, font_size_pt=None, color_rgb=None, font_name=None):
    """
    글자 모양 설정
    한글 2018 COM API를 사용하여 현재 커서 위치의 글자 모양 변경
    """
    try:
        act = hwp.CreateAction("CharShape")
        pset = act.CreateSet()
        act.GetDefault(pset)

        # 굵게
        if bold is not None:
            pset.SetItem("Bold", 1 if bold else 0)

        # 기울임
        if italic is not None:
            pset.SetItem("Italic", 1 if italic else 0)

        # 밑줄
        if underline is not None:
            pset.SetItem("Underline", 1 if underline else 0)

        # 글자 크기 (포인트 -> hwp 단위: pt * 100)
        if font_size_pt is not None:
            pset.SetItem("Height", int(font_size_pt * 100))

        # 글자 색상 (RGB -> BGR)
        if color_rgb is not None:
            r, g, b = color_rgb
            bgr = (b << 16) | (g << 8) | r
            pset.SetItem("TextColor", bgr)

        # 글꼴 이름
        if font_name is not None:
            pset.SetItem("FaceNameHangul", font_name)
            pset.SetItem("FaceNameLatin", font_name)
            pset.SetItem("FaceNameHanja", font_name)
            pset.SetItem("FaceNameJapanese", font_name)
            pset.SetItem("FaceNameOther", font_name)
            pset.SetItem("FaceNameSymbol", font_name)
            pset.SetItem("FaceNameUser", font_name)

        act.Execute(pset)
        logger.debug(f"글자 서식 적용: bold={bold}, italic={italic}, size={font_size_pt}")

    except Exception as e:
        logger.warning(f"글자 서식 적용 실패: {e}")

def _hwp_set_alignment(hwp, alignment):
    """
    문단 정렬 설정 (Run 액션 방식)
    """
    try:
        align_commands = {
            WD_ALIGN_PARAGRAPH.LEFT: "ParagraphShapeAlignLeft",
            WD_ALIGN_PARAGRAPH.CENTER: "ParagraphShapeAlignCenter",
            WD_ALIGN_PARAGRAPH.RIGHT: "ParagraphShapeAlignRight",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "ParagraphShapeAlignJustify",
            0: "ParagraphShapeAlignLeft",
            1: "ParagraphShapeAlignCenter",
            2: "ParagraphShapeAlignRight",
            3: "ParagraphShapeAlignJustify"
        }

        cmd = align_commands.get(alignment)
        if cmd:
            hwp.Run(cmd)
            logger.debug(f"정렬 적용: {cmd}")
    except Exception as e:
        logger.warning(f"정렬 적용 실패: {e}")

def _hwp_set_para_shape(hwp, alignment=None, left_margin=None, right_margin=None, first_line_indent=None,
                        line_spacing=None, space_before=None, space_after=None):
    """
    문단 모양 설정
    한글 2018 COM API를 사용하여 현재 문단의 모양 변경
    """
    try:
        act = hwp.CreateAction("ParagraphShape")
        pset = act.CreateSet()
        act.GetDefault(pset)

        # 왼쪽 여백 (포인트 -> hwp 단위: pt * 100)
        if left_margin is not None:
            pset.SetItem("LeftMargin", int(left_margin * 100))

        # 오른쪽 여백
        if right_margin is not None:
            pset.SetItem("RightMargin", int(right_margin * 100))

        # 첫줄 들여쓰기
        if first_line_indent is not None:
            pset.SetItem("IndentFirst", int(first_line_indent * 100))

        # 줄 간격 (0: 고정값, 1: 배수)
        if line_spacing is not None:
            pset.SetItem("LineSpacingType", 1)  # 배수로 설정
            pset.SetItem("LineSpacing", int(line_spacing * 100))

        # 문단 위 간격
        if space_before is not None:
            pset.SetItem("SpaceBefore", int(space_before * 100))

        # 문단 아래 간격
        if space_after is not None:
            pset.SetItem("SpaceAfter", int(space_after * 100))

        act.Execute(pset)
        logger.debug(f"문단 서식 적용: left={left_margin}, first={first_line_indent}")

    except Exception as e:
        logger.warning(f"문단 서식 적용 실패: {e}")

def _hwp_insert_table(hwp, rows, cols):
    """
    표 삽입
    """
    try:
        act = hwp.CreateAction("TableCreate")
        pset = act.CreateSet()
        act.GetDefault(pset)

        pset.SetItem("Rows", rows)
        pset.SetItem("Cols", cols)
        pset.SetItem("WidthType", 0)  # 0: 본문과 같게
        pset.SetItem("HeightType", 0)  # 0: 자동
        pset.SetItem("CreateItemArray", "TableProperties")

        act.Execute(pset)
        logger.info(f"표 생성: {rows}행 x {cols}열")
        return True

    except Exception as e:
        logger.error(f"표 생성 실패: {e}")
        return False

def _hwp_move_to_next_cell(hwp):
    """다음 셀로 이동 (Tab 키)"""
    try:
        hwp.HAction.Run("TableRightCell")
        return True
    except Exception as e:
        logger.error(f"다음 셀 이동 실패: {e}")
        return False

def _hwp_move_to_table_start(hwp):
    """표의 첫 번째 셀로 이동"""
    try:
        hwp.HAction.Run("TableLeftTopCell")
        return True
    except Exception as e:
        logger.error(f"표 시작으로 이동 실패: {e}")
        return False

def _hwp_set_cell_background(hwp, rgb_color):
    """
    셀 배경색 설정
    rgb_color: (r, g, b) 튜플
    """
    try:
        r, g, b = rgb_color
        bgr = (b << 16) | (g << 8) | r  # RGB -> BGR 변환

        act = hwp.CreateAction("CellBorderFill")
        pset = act.CreateSet()
        act.GetDefault(pset)

        # 배경색 설정
        pset.SetItem("FillAttr", 1)  # 단색 채우기
        pset.SetItem("FaceColor", bgr)

        act.Execute(pset)
        logger.debug(f"셀 배경색 설정: RGB{rgb_color}")
        return True
    except Exception as e:
        logger.warning(f"셀 배경색 설정 실패: {e}")
        return False

def _hwp_merge_cells(hwp, start_row, start_col, end_row, end_col):
    """
    셀 병합
    start_row, start_col: 시작 셀 (0-based)
    end_row, end_col: 끝 셀 (0-based)
    """
    try:
        # HAction 방식으로 셀 블록 선택
        hwp.HAction.GetDefault("TableCellBlock", hwp.HParameterSet.HTableCellBlock.HSet)
        hwp.HParameterSet.HTableCellBlock.StartRow = start_row
        hwp.HParameterSet.HTableCellBlock.StartCol = start_col
        hwp.HParameterSet.HTableCellBlock.EndRow = end_row
        hwp.HParameterSet.HTableCellBlock.EndCol = end_col
        hwp.HAction.Execute("TableCellBlock", hwp.HParameterSet.HTableCellBlock.HSet)

        # 셀 병합 실행
        hwp.HAction.Run("TableMergeCell")

        logger.debug(f"셀 병합: ({start_row},{start_col}) ~ ({end_row},{end_col})")
        return True
    except Exception as e:
        logger.warning(f"셀 병합 실패: {e}")
        return False

# -----------------------------------------------------------------------------
# DOCX 파싱 및 변환 함수들
# -----------------------------------------------------------------------------

def _get_cell_background_color(cell):
    """
    DOCX 셀의 배경색 추출
    Returns: (r, g, b) 튜플 또는 None
    """
    try:
        from docx.oxml.ns import qn

        # <w:shd> 요소에서 배경색 추출
        tcPr = cell._element.tcPr
        if tcPr is None:
            return None

        shd = tcPr.find(qn('w:shd'))
        if shd is None:
            return None

        fill = shd.get(qn('w:fill'))
        if fill is None or fill == 'auto':
            return None

        # 16진수 색상 코드를 RGB로 변환
        if len(fill) == 6:
            r = int(fill[0:2], 16)
            g = int(fill[2:4], 16)
            b = int(fill[4:6], 16)
            return (r, g, b)

    except Exception as e:
        logger.debug(f"배경색 추출 실패: {e}")

    return None

def _get_cell_span_info(cell):
    """
    DOCX 셀의 병합 정보 추출
    Returns: (grid_span, v_merge) - grid_span은 가로 병합 수, v_merge는 'continue' 또는 None
    """
    try:
        from docx.oxml.ns import qn

        tcPr = cell._element.tcPr
        if tcPr is None:
            return (1, None)

        # 가로 병합 (gridSpan)
        grid_span_elem = tcPr.find(qn('w:gridSpan'))
        grid_span = 1
        if grid_span_elem is not None:
            val = grid_span_elem.get(qn('w:val'))
            if val:
                grid_span = int(val)

        # 세로 병합 (vMerge)
        v_merge_elem = tcPr.find(qn('w:vMerge'))
        v_merge = None
        if v_merge_elem is not None:
            val = v_merge_elem.get(qn('w:val'))
            # val이 없으면 'continue', 있으면 'restart' (시작 셀)
            v_merge = 'continue' if val is None else 'restart'

        return (grid_span, v_merge)

    except Exception as e:
        logger.debug(f"셀 병합 정보 추출 실패: {e}")
        return (1, None)

def _convert_paragraph(hwp, para):
    """
    DOCX 문단을 HWP로 변환
    정렬 + Run별 텍스트 삽입 + 문단 나누기 후 들여쓰기 적용
    """
    try:
        # 1. 정렬 먼저 적용
        para_format = para.paragraph_format
        alignment = para.alignment if para.alignment else WD_ALIGN_PARAGRAPH.LEFT

        # 정렬 적용 (Run 액션 방식)
        if alignment is not None:
            _hwp_set_alignment(hwp, alignment)

        # 2. Run별 텍스트 삽입
        # Run이 없으면 단순 텍스트 삽입
        if not para.runs or len(para.runs) == 0:
            if para.text.strip():
                _hwp_insert_text(hwp, para.text)
        else:
            # Run별로 처리
            for run in para.runs:
                if not run.text:
                    continue

                # 글자 서식 추출
                bold = run.bold
                italic = run.italic
                underline = run.underline

                # 폰트 크기
                font_size = None
                if run.font.size:
                    font_size = run.font.size.pt

                # 폰트 색상
                color_rgb = None
                if run.font.color and run.font.color.rgb:
                    color_rgb = (run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])

                # 폰트 이름
                font_name = run.font.name

                # 글자 서식을 먼저 설정 (다음에 삽입될 텍스트에 적용됨)
                if bold is not None or italic is not None or underline is not None or font_size or color_rgb or font_name:
                    _hwp_set_char_shape(
                        hwp,
                        bold=bold,
                        italic=italic,
                        underline=underline,
                        font_size_pt=font_size,
                        color_rgb=color_rgb,
                        font_name=font_name
                    )

                # 텍스트 삽입
                _hwp_insert_text(hwp, run.text)

        # 3. 들여쓰기 적용 (문단 나누기 전에)
        left_indent = para_format.left_indent.pt if para_format.left_indent else 0
        right_indent = para_format.right_indent.pt if para_format.right_indent else 0
        first_line_indent = para_format.first_line_indent.pt if para_format.first_line_indent else 0

        # 문단 간격
        space_before = para_format.space_before.pt if para_format.space_before else 0
        space_after = para_format.space_after.pt if para_format.space_after else 0

        # 들여쓰기/간격 적용 (기본값이 아닐 때만)
        if (left_indent != 0 or right_indent != 0 or first_line_indent != 0 or
            space_before != 0 or space_after != 0):
            _hwp_set_para_shape(
                hwp,
                left_margin=left_indent,
                right_margin=right_indent,
                first_line_indent=first_line_indent,
                space_before=space_before,
                space_after=space_after
            )

        # 4. 문단 나누기
        _hwp_break_para(hwp)

    except Exception as e:
        logger.error(f"문단 변환 실패: {e}")
        import traceback
        logger.error(traceback.format_exc())

def _analyze_table_merges(table):
    """
    DOCX 표의 병합 정보 분석
    Returns: merge_info dict
        - merge_map: {(row, col): {'h_span': int, 'v_span': int, 'skip': bool}}
        - merge_list: [(start_row, start_col, end_row, end_col), ...]
    """
    rows = len(table.rows)
    cols = len(table.columns)

    merge_map = {}
    v_merge_tracker = {}  # {col: (start_row, current_row)}

    # 1단계: 모든 셀 순회하며 병합 정보 수집
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            grid_span, v_merge = _get_cell_span_info(cell)

            # 가로 병합 정보
            h_span = grid_span

            # 세로 병합 정보
            v_span = 1
            skip = False

            if v_merge == 'restart':
                # 세로 병합 시작
                v_merge_tracker[col_idx] = row_idx
            elif v_merge == 'continue':
                # 세로 병합 계속 (이 셀은 병합된 셀)
                skip = True
                if col_idx in v_merge_tracker:
                    start_row = v_merge_tracker[col_idx]
                    merge_map[(start_row, col_idx)] = merge_map.get((start_row, col_idx), {'h_span': 1, 'v_span': 1, 'skip': False})
                    merge_map[(start_row, col_idx)]['v_span'] = row_idx - start_row + 1
            else:
                # 병합 없음 또는 종료
                if col_idx in v_merge_tracker:
                    del v_merge_tracker[col_idx]

            merge_map[(row_idx, col_idx)] = {
                'h_span': h_span,
                'v_span': v_span,
                'skip': skip
            }

    # 2단계: 병합 목록 생성 (HWP에서 병합할 셀 범위)
    merge_list = []
    for (row, col), info in merge_map.items():
        if not info['skip']:  # 병합 시작 셀만
            h_span = info['h_span']
            v_span = info['v_span']

            if h_span > 1 or v_span > 1:
                end_row = row + v_span - 1
                end_col = col + h_span - 1
                merge_list.append((row, col, end_row, end_col))
                logger.debug(f"병합 감지: ({row},{col}) ~ ({end_row},{end_col})")

    return {
        'merge_map': merge_map,
        'merge_list': merge_list
    }

def _convert_table(hwp, table):
    """
    DOCX 표를 HWP로 변환
    """
    try:
        rows = len(table.rows)
        cols = len(table.columns)

        logger.info(f"표 변환 시작: {rows}행 x {cols}열")

        # 병합 정보 분석
        merge_info = _analyze_table_merges(table)
        merge_map = merge_info['merge_map']
        merge_list = merge_info['merge_list']

        logger.info(f"병합 정보: {len(merge_list)}개 병합 영역")

        # 표 생성 (생성 후 커서는 자동으로 첫 번째 셀에 위치)
        if not _hwp_insert_table(hwp, rows, cols):
            return

        # 표의 첫 번째 셀로 이동 확인
        _hwp_move_to_table_start(hwp)

        # 각 셀에 내용 채우기
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # 병합된 셀(skip=True)은 건너뛰기
                cell_info = merge_map.get((row_idx, col_idx), {'skip': False})
                if cell_info['skip']:
                    logger.debug(f"셀 ({row_idx}, {col_idx}) 병합된 셀 건너뛰기")
                    continue

                # 현재 셀에 내용 삽입
                try:
                    # 셀 내용 삽입 (문단들)
                    for para_idx, para in enumerate(cell.paragraphs):
                        if para.text.strip():
                            # 정렬 적용 (문단의 정렬)
                            para_alignment = para.alignment if para.alignment else WD_ALIGN_PARAGRAPH.LEFT
                            _hwp_set_alignment(hwp, para_alignment)
                            # 첫 번째 문단이 아니면 줄바꿈
                            if para_idx > 0:
                                _hwp_break_para(hwp)

                            # Run별 텍스트 삽입 (문단 나누기 제외)
                            for run in para.runs:
                                if not run.text:
                                    continue

                                # 글자 서식 추출
                                bold = run.bold
                                italic = run.italic
                                underline = run.underline

                                # 폰트 크기
                                font_size = None
                                if run.font.size:
                                    font_size = run.font.size.pt

                                # 폰트 색상
                                color_rgb = None
                                if run.font.color and run.font.color.rgb:
                                    color_rgb = (run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])

                                # 폰트 이름
                                font_name = run.font.name

                                # 글자 서식 설정
                                if bold is not None or italic is not None or underline is not None or font_size or color_rgb or font_name:
                                    _hwp_set_char_shape(
                                        hwp,
                                        bold=bold,
                                        italic=italic,
                                        underline=underline,
                                        font_size_pt=font_size,
                                        color_rgb=color_rgb,
                                        font_name=font_name
                                    )

                                # 텍스트 삽입
                                _hwp_insert_text(hwp, run.text)

                    logger.debug(f"셀 ({row_idx}, {col_idx}) 내용 입력 완료")

                except Exception as e:
                    logger.warning(f"셀 ({row_idx}, {col_idx}) 처리 중 오류: {e}")

                # 다음 셀로 이동 (Tab 키) - 마지막 셀이 아닌 경우만
                if not (row_idx == rows - 1 and col_idx == cols - 1):
                    _hwp_move_to_next_cell(hwp)

        # 셀 병합 실행
        if merge_list:
            logger.info(f"셀 병합 시작: {len(merge_list)}개 병합 영역")
            for start_row, start_col, end_row, end_col in merge_list:
                try:
                    _hwp_merge_cells(hwp, start_row, start_col, end_row, end_col)
                    logger.debug(f"병합 완료: ({start_row},{start_col}) ~ ({end_row},{end_col})")
                except Exception as e:
                    logger.warning(f"병합 실패 ({start_row},{start_col}) ~ ({end_row},{end_col}): {e}")

        # 표 밖으로 나가기
        hwp.HAction.Run("TableRightTopCell")
        hwp.HAction.Run("BreakPara")

        logger.info("표 변환 완료")

    except Exception as e:
        logger.error(f"표 변환 실패: {e}")
        import traceback
        logger.error(traceback.format_exc())

def _convert_document_body(hwp, doc):
    """
    DOCX 본문을 HWP로 변환
    """
    try:
        total_elements = len(doc.element.body)
        logger.info(f"본문 요소 {total_elements}개 변환 시작")

        for idx, element in enumerate(doc.element.body):
            # 문단인 경우
            if isinstance(element, CT_P):
                para = None
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break

                if para:
                    logger.debug(f"[{idx+1}/{total_elements}] 문단 변환: {para.text[:30]}...")
                    _convert_paragraph(hwp, para)

            # 표인 경우
            elif isinstance(element, CT_Tbl):
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break

                if table:
                    logger.debug(f"[{idx+1}/{total_elements}] 표 변환")
                    _convert_table(hwp, table)

            if (idx + 1) % 10 == 0:
                logger.info(f"진행중... {idx+1}/{total_elements}")

        logger.info("본문 변환 완료")

    except Exception as e:
        logger.error(f"본문 변환 실패: {e}")
        import traceback
        logger.error(traceback.format_exc())

# -----------------------------------------------------------------------------
# 메인 변환 함수
# -----------------------------------------------------------------------------

def perform_conversion(docx_path: str, hwp_path: str) -> bool:
    """
    한글 2018 COM API를 사용하여 DOCX를 HWP로 완전 수동 변환
    """
    hwp = None
    try:
        # COM 초기화
        pythoncom.CoInitialize()

        logger.info("=" * 70)
        logger.info("한글 2018 COM API 수동 변환 시작")
        logger.info("=" * 70)

        # HWP 오브젝트 생성
        hwp = win32com.client.Dispatch("HWPFrame.HwpObject")
        hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        hwp.SetMessageBoxMode(0x00010000)  # 메시지 박스 표시 안함

        logger.info(f"한글 버전: {hwp.Version}")

        # 절대 경로 변환
        abs_docx_path = os.path.abspath(docx_path)
        abs_hwp_path = os.path.abspath(hwp_path)

        logger.info(f"입력: {abs_docx_path}")
        logger.info(f"출력: {abs_hwp_path}")

        # 기존 파일 삭제
        if os.path.exists(abs_hwp_path):
            try:
                os.remove(abs_hwp_path)
                logger.info("기존 HWP 파일 삭제 완료")
            except Exception as e:
                logger.warning(f"기존 파일 삭제 실패: {e}")

        # DOCX 문서 로드
        logger.info("DOCX 파일 로드 중...")
        doc = Document(abs_docx_path)
        logger.info(f"문단 수: {len(doc.paragraphs)}, 표 수: {len(doc.tables)}")

        # 빈 HWP 문서 생성
        logger.info("빈 HWP 문서 생성...")
        hwp.HAction.Run("FileNew")

        # 본문 변환
        logger.info("본문 변환 시작...")
        _convert_document_body(hwp, doc)

        # 파일 저장
        logger.info("HWP 파일 저장 중...")
        hwp.SaveAs(abs_hwp_path, "HWP", "")

        # 저장 확인
        if os.path.exists(abs_hwp_path):
            file_size = os.path.getsize(abs_hwp_path)
            logger.info(f"✓ 변환 성공! 파일 크기: {file_size:,} bytes")
            logger.info(f"✓ 저장 위치: {abs_hwp_path}")
            return True
        else:
            logger.error("✗ 파일이 생성되지 않았습니다")
            return False

    except Exception as e:
        logger.error(f"✗ 변환 실패: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return False

    finally:
        # HWP 종료
        if hwp:
            try:
                hwp.Quit()
                logger.info("한글 프로세스 종료")
            except:
                pass
            del hwp

        # COM 해제
        pythoncom.CoUninitialize()
        logger.info("=" * 70)

# -----------------------------------------------------------------------------
# 래퍼 클래스 (main.py와 호환성 유지)
# -----------------------------------------------------------------------------

class HWPConverter:
    """HWP 변환기 클래스"""

    def convert_docx_to_hwp(self, docx_path: str, hwp_path: str) -> bool:
        """DOCX를 HWP로 변환"""
        logger.info(f"변환 요청: {os.path.basename(docx_path)} -> {os.path.basename(hwp_path)}")
        return perform_conversion(docx_path, hwp_path)

    def check_hwp_installed(self):
        """한글 설치 여부 확인"""
        hwp = None
        try:
            pythoncom.CoInitialize()
            hwp = win32com.client.Dispatch("HWPFrame.HwpObject")

            version = hwp.Version
            version_str = str(version)

            logger.info(f"한글 설치 확인: {version_str}")
            return True, version_str

        except Exception as e:
            logger.error(f"한글 확인 실패: {e}")
            return False, "Not Installed"

        finally:
            if hwp:
                try:
                    hwp.Quit()
                except:
                    pass
                del hwp
            pythoncom.CoUninitialize()
