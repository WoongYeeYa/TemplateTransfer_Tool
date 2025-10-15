"""
문서 분석 서비스
Word 및 HWP 문서 분석을 위한 MCP 통합
"""
import os
from typing import Dict, Any, List
from docx import Document as DocxDocument


class DocumentAnalyzer:
    """문서 분석 서비스"""

    async def analyze_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        문서 분석 메인 함수

        Args:
            file_path: 분석할 파일 경로
            file_type: 파일 타입 ('.docx', '.hwp', '.doc')

        Returns:
            분석 결과 딕셔너리
        """
        if file_type in ['.docx', '.doc']:
            return await self.analyze_word_document(file_path)
        elif file_type == '.hwp':
            return await self.analyze_hwp_document(file_path)
        else:
            raise ValueError(f"지원하지 않는 파일 형식: {file_type}")

    async def analyze_word_document(self, file_path: str) -> Dict[str, Any]:
        """
        Word 문서 분석
        python-docx를 사용한 기본 분석
        추후 Office-Word-MCP로 확장 가능
        """
        try:
            doc = DocxDocument(file_path)

            # 기본 정보 추출
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })

            # 표 정보 추출
            tables = []
            for table in doc.tables:
                table_data = {
                    'rows': len(table.rows),
                    'cols': len(table.columns) if table.rows else 0,
                    'cells': []
                }

                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data['cells'].append(row_data)

                tables.append(table_data)

            # 필드 자동 인식 (간단한 패턴 매칭)
            fields = self._identify_fields(paragraphs, tables)

            return {
                'document_type': 'word',
                'paragraphs': paragraphs,
                'tables': tables,
                'fields': fields,
                'summary': {
                    'paragraph_count': len(paragraphs),
                    'table_count': len(tables),
                    'field_count': len(fields)
                }
            }

        except Exception as e:
            return {
                'error': f"Word 문서 분석 실패: {str(e)}",
                'document_type': 'word',
                'paragraphs': [],
                'tables': [],
                'fields': []
            }

    async def analyze_hwp_document(self, file_path: str) -> Dict[str, Any]:
        """
        HWP 문서 분석
        추후 hwpx MCP 통합
        현재는 기본 구조만 반환
        """
        # TODO: hwpx MCP 통합
        return {
            'document_type': 'hwp',
            'paragraphs': [],
            'tables': [],
            'fields': [],
            'summary': {
                'paragraph_count': 0,
                'table_count': 0,
                'field_count': 0
            },
            'note': 'HWP 분석 기능은 추후 hwpx MCP 통합 예정'
        }

    def _identify_fields(self, paragraphs: List[Dict], tables: List[Dict]) -> List[Dict]:
        """
        AI 기반 필드 자동 인식 (간단한 패턴 매칭)
        추후 AI 모델로 확장 가능
        """
        fields = []
        field_patterns = {
            '이름': ['이름', '성명', 'name', '신청자'],
            '날짜': ['날짜', '일자', 'date', '작성일'],
            '부서': ['부서', '소속', 'department', '팀'],
            '전화번호': ['전화', '연락처', 'phone', 'tel'],
            '이메일': ['이메일', 'email', 'e-mail'],
            '주소': ['주소', 'address', '거주지'],
        }

        # 문단에서 필드 패턴 찾기
        for para in paragraphs:
            text = para['text'].lower()
            for field_name, patterns in field_patterns.items():
                for pattern in patterns:
                    if pattern in text and ':' in text:
                        fields.append({
                            'id': f'field_{len(fields) + 1}',
                            'name': field_name,
                            'type': 'text',
                            'source': 'paragraph',
                            'original_text': para['text']
                        })
                        break

        # 표에서 필드 찾기
        for table_idx, table in enumerate(tables):
            for row_idx, row in enumerate(table['cells']):
                for col_idx, cell in enumerate(row):
                    cell_lower = cell.lower()
                    for field_name, patterns in field_patterns.items():
                        for pattern in patterns:
                            if pattern in cell_lower:
                                fields.append({
                                    'id': f'field_{len(fields) + 1}',
                                    'name': field_name,
                                    'type': 'text',
                                    'source': 'table',
                                    'table_index': table_idx,
                                    'row': row_idx,
                                    'col': col_idx
                                })
                                break

        return fields
