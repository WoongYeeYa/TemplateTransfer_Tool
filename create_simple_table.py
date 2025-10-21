"""
간단한 표가 포함된 DOCX 파일 생성 (셀 병합 없음)
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_simple_table():
    """간단한 3x3 표 생성"""
    doc = Document()

    # 제목
    title = doc.add_paragraph()
    title_run = title.add_run("간단한 표 테스트")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 3x3 표 생성
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    # 첫 번째 행 - 헤더
    headers = ["이름", "나이", "직업"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(12)

    # 두 번째 행
    row1_data = ["홍길동", "30", "개발자"]
    for idx, data in enumerate(row1_data):
        cell = table.rows[1].cells[idx]
        cell.text = data
        para = cell.paragraphs[0]
        run = para.runs[0]
        run.font.size = Pt(11)

    # 세 번째 행
    row2_data = ["김철수", "25", "디자이너"]
    for idx, data in enumerate(row2_data):
        cell = table.rows[2].cells[idx]
        cell.text = data
        para = cell.paragraphs[0]
        run = para.runs[0]
        run.font.size = Pt(11)

    # 파일 저장
    filename = "test_simple_table.docx"
    doc.save(filename)
    print(f"[OK] 간단한 표 테스트 파일 생성: {filename}")

    return filename

if __name__ == "__main__":
    create_simple_table()
