"""
복잡한 표가 포함된 DOCX 파일 생성
- 셀 병합 (가로/세로)
- 셀 배경색
- 표 안의 정렬
- 폰트 색깔
- 폰트 크기
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill):
    """셀 배경색 설정"""
    try:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill)
        cell._element.get_or_add_tcPr().append(shading_elm)
    except Exception as e:
        print(f"배경색 설정 실패: {e}")

def create_complex_table_docx():
    """복잡한 표가 포함된 DOCX 생성"""
    doc = Document()

    # 제목
    title = doc.add_paragraph()
    title_run = title.add_run("복잡한 표 변환 테스트")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 128)  # 네이비
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # === 테스트 1: 셀 병합 + 배경색 + 정렬 ===
    heading1 = doc.add_paragraph()
    heading1_run = heading1.add_run("1. 셀 병합 + 배경색 + 정렬 테스트")
    heading1_run.font.size = Pt(14)
    heading1_run.font.bold = True
    heading1_run.font.color.rgb = RGBColor(255, 0, 0)

    # 5x5 표 생성
    table1 = doc.add_table(rows=5, cols=5)
    table1.style = 'Table Grid'

    # 첫 번째 행 - 제목 행 (가로 병합)
    cell = table1.rows[0].cells[0]
    cell.merge(table1.rows[0].cells[4])  # 5개 셀 병합
    cell.text = "제목 행 (5개 셀 병합)"
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)  # 흰색
    set_cell_background(cell, "4472C4")  # 파란 배경

    # 두 번째 행 - 헤더 (각각 다른 배경색)
    headers = ["이름", "나이", "직업", "지역", "비고"]
    colors = ["FFD966", "C6E0B4", "F4B084", "B4C7E7", "D5A6BD"]
    for idx, (header, color) in enumerate(zip(headers, colors)):
        cell = table1.rows[1].cells[idx]
        cell.text = header
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(12)
        set_cell_background(cell, color)

    # 세 번째 행 - 일반 데이터 (다양한 정렬)
    data_row1 = ["홍길동", "30", "개발자", "서울", "우수"]
    alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.JUSTIFY]
    font_colors = [
        RGBColor(0, 0, 0),      # 검정
        RGBColor(255, 0, 0),    # 빨강
        RGBColor(0, 128, 0),    # 초록
        RGBColor(0, 0, 255),    # 파랑
        RGBColor(128, 0, 128)   # 보라
    ]

    for idx, (data, align, color) in enumerate(zip(data_row1, alignments, font_colors)):
        cell = table1.rows[2].cells[idx]
        cell.text = data
        para = cell.paragraphs[0]
        para.alignment = align
        run = para.runs[0]
        run.font.color.rgb = color
        run.font.size = Pt(11)

    # 네 번째 행 - 세로 병합
    cell = table1.rows[3].cells[0]
    cell.merge(table1.rows[4].cells[0])  # 2행 병합
    cell.text = "세로\n병합\n셀"
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_background(cell, "FF6B6B")  # 빨간 배경

    # 네 번째 행 나머지 - 작은 폰트
    for idx in range(1, 5):
        cell = table1.rows[3].cells[idx]
        cell.text = f"작은 텍스트 {idx}"
        para = cell.paragraphs[0]
        run = para.runs[0]
        run.font.size = Pt(8)

    # 다섯 번째 행 - 큰 폰트
    for idx in range(1, 5):
        cell = table1.rows[4].cells[idx]
        cell.text = f"큰 글자"
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 100, 0)  # 주황

    doc.add_paragraph()

    # === 테스트 2: 복잡한 셀 병합 패턴 ===
    heading2 = doc.add_paragraph()
    heading2_run = heading2.add_run("2. 복잡한 셀 병합 패턴")
    heading2_run.font.size = Pt(14)
    heading2_run.font.bold = True
    heading2_run.font.color.rgb = RGBColor(0, 128, 0)

    # 4x4 표
    table2 = doc.add_table(rows=4, cols=4)
    table2.style = 'Table Grid'

    # L자 모양 병합
    # 첫 번째 행 첫 2개 셀 병합
    cell = table2.rows[0].cells[0]
    cell.merge(table2.rows[0].cells[1])
    cell.text = "가로 병합"
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    set_cell_background(cell, "A8DADC")

    # 첫 번째 열 2-3행 병합
    cell = table2.rows[1].cells[0]
    cell.merge(table2.rows[2].cells[0])
    cell.text = "세로\n병합"
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    set_cell_background(cell, "F1FAEE")

    # 2x2 블록 병합 (오른쪽 하단)
    cell = table2.rows[2].cells[2]
    cell.merge(table2.rows[3].cells[3])
    cell.text = "2x2 병합"
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_background(cell, "457B9D")

    # 나머지 셀들
    table2.rows[0].cells[2].text = "A"
    table2.rows[0].cells[3].text = "B"
    table2.rows[1].cells[1].text = "C"
    table2.rows[1].cells[2].text = "D"
    table2.rows[1].cells[3].text = "E"
    table2.rows[2].cells[1].text = "F"
    table2.rows[3].cells[0].text = "G"
    table2.rows[3].cells[1].text = "H"

    doc.add_paragraph()

    # === 테스트 3: 다양한 폰트 크기와 색상 ===
    heading3 = doc.add_paragraph()
    heading3_run = heading3.add_run("3. 다양한 폰트 크기와 색상")
    heading3_run.font.size = Pt(14)
    heading3_run.font.bold = True
    heading3_run.font.color.rgb = RGBColor(255, 0, 255)

    # 3x6 표
    table3 = doc.add_table(rows=3, cols=6)
    table3.style = 'Table Grid'

    font_sizes = [8, 10, 12, 14, 16, 18]
    gradient_colors = [
        RGBColor(255, 0, 0),    # 빨강
        RGBColor(255, 128, 0),  # 주황
        RGBColor(255, 255, 0),  # 노랑
        RGBColor(0, 255, 0),    # 초록
        RGBColor(0, 128, 255),  # 하늘
        RGBColor(128, 0, 255)   # 보라
    ]

    # 첫 번째 행 - 폰트 크기
    for idx, size in enumerate(font_sizes):
        cell = table3.rows[0].cells[idx]
        cell.text = f"{size}pt"
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.size = Pt(size)
        run.font.bold = True

    # 두 번째 행 - 색상
    for idx, color in enumerate(gradient_colors):
        cell = table3.rows[1].cells[idx]
        cell.text = "색상"
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.size = Pt(12)
        run.font.color.rgb = color
        run.font.bold = True

    # 세 번째 행 - 배경색 그라데이션
    bg_colors = ["FFCCCC", "FFDDAA", "FFFFCC", "CCFFCC", "CCDDFF", "DDCCFF"]
    for idx, bg_color in enumerate(bg_colors):
        cell = table3.rows[2].cells[idx]
        cell.text = "배경"
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        set_cell_background(cell, bg_color)

    # 파일 저장
    filename = "test_complex_table.docx"
    doc.save(filename)
    print(f"[OK] 복잡한 표 테스트 파일 생성 완료: {filename}")
    print()
    print("테스트 포함 내용:")
    print("  1. 셀 병합 (가로 5개, 세로 2개, 2x2 블록)")
    print("  2. 셀 배경색 (10가지 이상)")
    print("  3. 표 안의 정렬 (왼쪽, 가운데, 오른쪽, 양쪽)")
    print("  4. 폰트 색깔 (빨강, 초록, 파랑, 보라, 주황, 노랑 등)")
    print("  5. 폰트 크기 (8pt ~ 18pt)")
    print("  6. 복잡한 병합 패턴 (L자, 2x2 블록 등)")

    return filename

if __name__ == "__main__":
    create_complex_table_docx()
